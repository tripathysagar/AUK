

import parse_json
import pie #for generating pie chat
from docx import Document
from docx.shared import Inches


file = 'demo registratio failed.postman_test_run.json'
#file = 'demo registratio pass.postman_test_run.json'
input_for_chart  = parse_json.parse(file)
pie.main(input_for_chart, file)

document = Document(file + '_report.docx')
p = document.add_paragraph()
r = p.add_run()
r.add_picture('chart.jpg', width=Inches(5.8))
document.add_paragraph('Inner and outer circle represents requests and test results respectively')
document.save(file + '_report.docx')
