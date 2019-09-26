import json
from pprint import pprint

from docx import Document
from docx.shared import Inches
document = Document()


def add_line( string):
    p = document.add_paragraph('',style='List Bullet')

    p.add_run(string)

#s += 'sucess : ' + str(sucessCount) + '\nfailure : ' + str(failCount)
#f.write(string)

#returns two value after parsing result of a collection
#result_for_pie_chat = {'login': [1, 1], 'getUserInfo': [2, 4], 'updation': [7, 3]} for drawing the pie pie_chart
#and dictionary of suces and fail count
#it also saves output of each request in "file + 'report.pdf'" file
def parse(file):
    with open(file)    as data_file:
        data = json.load(data_file)


    document.add_heading(file, 0)
    #f = open( 'report.txt', 'w')

    #pprint(len(data["results"]))


    result_for_pie_chat = {} #for calculating the sucesas and failure rate of the


    for i in range(len(data["results"])):
        sucessCount = 0
        failCount = 0

        dic = data["results"][i]
        string = ''

        document.add_heading('Request name : \t' + dic["name"] + '\n', level=1)


        add_line('id :\t' + dic["id"] + '\n')

        add_line( 'url :\t' + dic["url"] + '\n')
        add_line( 'code :\t' + str(dic["responseCode"]["code"]) + '\n')
        add_line('totalRequestTime :\t' + dic["totalRequestTime"] + '\n')



        records = []
        req = dic["tests"].keys() #find the names of the requests

        string = 'tests and results\n\n'
        if len(req) > 0:    #if there are tests done in the request
            for i in req:
                temp = []
                if dic["testPassFailCounts"][i]["pass"] == 1: #if the test is sucess
                    sucessCount += 1
                    temp = [i, 'pass']
                    #string += i + ':\t' + 'pass\n'
                else:
                    failCount += 1
                    temp = [i, 'fail']
                    #string += i + ':\t' + 'fail\n'

                records.append(temp)
                #print(string)
            result_for_pie_chat[dic["name"]] = [sucessCount, failCount]    #saving the count of  requests result
            #print(records)


            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Test'
            hdr_cells[1].text = 'Result'

            for test, result in records:
                row_cells = table.add_row().cells
                row_cells[0].text = str(test)
                row_cells[1].text = result
        else:
            add_line('No Tests')

    document.add_page_break()

    document.save(file + '_report.docx')

    return result_for_pie_chat
#result , dic = parse(file)
#print(result)
#print(dic)
